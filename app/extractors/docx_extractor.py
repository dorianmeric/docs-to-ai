"""
DOCX Extractor - Extract text from Word documents.
"""

from pathlib import Path
from typing import List, Dict, Any
from docx import Document as DocxDocument


def extract_text_from_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """Extract text from Word document.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        List of dicts with 'page' (paragraph number) and 'text'
    """
    pages_data = []

    try:
        doc = DocxDocument(str(docx_path))
        
        paragraphs_per_page = 10
        current_text = []
        page_num = 1
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                current_text.append(text)
            
            if (i + 1) % paragraphs_per_page == 0 or i == len(doc.paragraphs) - 1:
                if current_text:
                    pages_data.append({
                        'page': page_num,
                        'text': '\n'.join(current_text),
                        'total_pages': -1
                    })
                    page_num += 1
                    current_text = []
        
    except Exception as e:
        print(f"Error processing Word document {docx_path}: {e}")
        return []
    
    return pages_data
