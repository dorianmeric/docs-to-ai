import pymupdf  # PyMuPDF -- Library for PDF processing
from pathlib import Path
from typing import List, Dict, Optional
import hashlib
import json
from docx import Document as DocxDocument
from config import CHUNK_SIZE, CHUNK_OVERLAP, DOC_CACHE_DIR, USE_FOLDER_AS_TOPIC, DEFAULT_TOPIC


class DocumentProcessor:
    """Handles document text extraction and chunking for PDFs and Word documents."""
    
    def __init__(self):
        self.cache_dir = DOC_CACHE_DIR
    
    def extract_topics_from_path(self, doc_path: Path, base_dir: Optional[Path] = None) -> List[str]:
        """
        Extract all topics from folder hierarchy.
        
        Args:
            doc_path: Path to the document file
            base_dir: Base directory for documents (to determine topic hierarchy)
            
        Returns:
            List of topic names from folder hierarchy
        """
        if not USE_FOLDER_AS_TOPIC:
            return [DEFAULT_TOPIC]
        
        # If no base directory, just use parent folder
        if not base_dir:
            parent_name = doc_path.parent.name
            if parent_name and parent_name not in ['/', '\\\\', '.']:
                return [parent_name]
            return [DEFAULT_TOPIC]
        
        base_path = Path(base_dir).resolve()
        doc_path_resolved = doc_path.resolve()
        
        # If document is directly in base directory
        if doc_path_resolved.parent == base_path:
            return [DEFAULT_TOPIC]
        
        # Get relative path from base to document's parent folder
        try:
            relative_path = doc_path_resolved.parent.relative_to(base_path)
        except ValueError:
            # Document is not under base directory
            return [doc_path.parent.name if doc_path.parent.name else DEFAULT_TOPIC]
        
        # Extract all folder names in the path as topics
        topics = [part for part in relative_path.parts if part and part not in ['.', '..']]
        
        return topics if topics else [DEFAULT_TOPIC]
    
    def extract_text_from_pdf(self, pdf_path: Path) -> List[Dict[str, any]]:
        """
        Extract text from PDF, maintaining page information.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of dicts with 'page' and 'text'
        """
        pages_data = []
        
        try:
            doc = pymupdf.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                pages_data.append({
                    'page': page_num + 1,
                    'text': text,
                    'total_pages': len(doc)
                })
            
            doc.close()
            
        except Exception as e:
            print(f"Error processing PDF {pdf_path}: {e}")
            return []
        
        return pages_data
    
    def extract_text_from_docx(self, docx_path: Path) -> List[Dict[str, any]]:
        """
        Extract text from Word document.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of dicts with 'page' (paragraph number) and 'text'
        """
        pages_data = []
        
        try:
            doc = DocxDocument(docx_path)
            
            # Combine paragraphs into chunks (simulating pages)
            # Group every ~10 paragraphs as a "page"
            paragraphs_per_page = 10
            current_text = []
            page_num = 1
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    current_text.append(text)
                
                # Create a "page" every N paragraphs or at the end
                if (i + 1) % paragraphs_per_page == 0 or i == len(doc.paragraphs) - 1:
                    if current_text:
                        pages_data.append({
                            'page': page_num,
                            'text': '\\n'.join(current_text),
                            'total_pages': -1  # Unknown for Word docs
                        })
                        page_num += 1
                        current_text = []
            
        except Exception as e:
            print(f"Error processing Word document {docx_path}: {e}")
            return []
        
        return pages_data
    
    def extract_text_from_document(self, doc_path: str, base_dir: Optional[str] = None) -> List[Dict[str, any]]:
        """
        Extract text from document (PDF or Word), maintaining structure.
        
        Args:
            doc_path: Path to the document file
            base_dir: Base directory for documents (to extract topics from folder structure)
            
        Returns:
            List of dicts with 'page', 'text', and 'metadata'
        """
        doc_path = Path(doc_path)
        base_path = Path(base_dir) if base_dir else None
        
        # Extract topics from folder hierarchy
        topics = self.extract_topics_from_path(doc_path, base_path)
        
        # Check cache
        cache_file = self._get_cache_path(doc_path)
        if cache_file.exists():
            cached_data = json.load(open(cache_file, 'r', encoding='utf-8'))
            # Update topics in cached data in case folder structure changed
            for page_data in cached_data:
                page_data['metadata']['topics'] = topics
            return cached_data
        
        # Determine file type and extract text
        extension = doc_path.suffix.lower()
        
        if extension == '.pdf':
            pages_data = self.extract_text_from_pdf(doc_path)
        elif extension in ['.docx', '.doc']:
            pages_data = self.extract_text_from_docx(doc_path)
        else:
            print(f"Unsupported file type: {extension}")
            return []
        
        # Add metadata to each page
        result = []
        for page_data in pages_data:
            result.append({
                'page': page_data['page'],
                'text': page_data['text'],
                'metadata': {
                    'filename': doc_path.name,
                    'filepath': str(doc_path),
                    'topics': topics,
                    'filetype': extension,
                    'total_pages': page_data['total_pages']
                }
            })
        
        # Cache the result
        self._cache_extracted_text(doc_path, result)
        
        return result
    
    def chunk_text(self, pages_data: List[Dict[str, any]]) -> List[Dict[str, any]]:
        """
        Split text into chunks with overlap.
        
        Args:
            pages_data: List of page data from extract_text_from_document
            
        Returns:
            List of chunks with metadata
        """
        chunks = []
        
        for page_data in pages_data:
            text = page_data['text']
            page_num = page_data['page']
            metadata = page_data['metadata']
            
            # Skip empty pages
            if not text.strip():
                continue
            
            # Split into chunks
            for i in range(0, len(text), CHUNK_SIZE - CHUNK_OVERLAP):
                chunk_text = text[i:i + CHUNK_SIZE]
                
                # Skip very small chunks
                if len(chunk_text.strip()) < 50:
                    continue
                
                # Include topics in chunk ID to ensure uniqueness
                topics_str = '-'.join(metadata['topics'])
                chunk_id = hashlib.md5(
                    f"{topics_str}-{metadata['filepath']}-{page_num}-{i}".encode()
                ).hexdigest()
                
                chunks.append({
                    'id': chunk_id,
                    'text': chunk_text,
                    'metadata': {
                        **metadata,
                        'page': page_num,
                        'chunk_start': i,
                        'chunk_end': i + len(chunk_text)
                    }
                })
        
        return chunks
    
    def process_document(self, doc_path: str, base_dir: Optional[str] = None) -> List[Dict[str, any]]:
        """
        Complete pipeline: extract text and chunk it.
        
        Args:
            doc_path: Path to document file (PDF or Word)
            base_dir: Base directory for documents (to extract topics from folder structure)
            
        Returns:
            List of text chunks with metadata
        """
        pages_data = self.extract_text_from_document(doc_path, base_dir)
        chunks = self.chunk_text(pages_data)
        return chunks
    
    def _get_cache_path(self, doc_path: Path) -> Path:
        """Generate cache file path for a document."""
        doc_hash = hashlib.md5(str(doc_path).encode()).hexdigest()
        return self.cache_dir / f"{doc_hash}.json"
    
    def _cache_extracted_text(self, doc_path: Path, pages_data: List[Dict]):
        """Cache extracted text to avoid reprocessing."""
        cache_file = self._get_cache_path(doc_path)
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump(pages_data, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    # Test the processor
    processor = DocumentProcessor()
    
    # Example usage
    test_doc = "test.pdf"  # Replace with actual document path
    if Path(test_doc).exists():
        chunks = processor.process_document(test_doc)
        print(f"Extracted {len(chunks)} chunks from {test_doc}")
        if chunks:
            print(f"\\nFirst chunk preview:")
            print(f"Topics: {chunks[0]['metadata']['topics']}")
            print(chunks[0]['text'][:200])
    else:
        print(f"Test file {test_doc} not found")
